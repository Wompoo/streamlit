import streamlit as st
import whisper
import tempfile
import os
from io import BytesIO
from docx import Document

st.title("Transcriptor de Audio a Texto con Whisper")

@st.cache_resource
def load_model():
    return whisper.load_model("small")

model = load_model()

uploaded_file = st.file_uploader("Sube un archivo de audio (.wav o .mp4)", type=["wav", "mp4"])

if uploaded_file is not None:
    st.success("Archivo recibido. Procesando...")

    with tempfile.NamedTemporaryFile(delete=False, suffix=uploaded_file.name[-4:]) as tmp_file:
        tmp_file.write(uploaded_file.read())
        tmp_path = tmp_file.name

    st.info("Transcribiendo, espera un momento...")
    try:
        result = model.transcribe(tmp_path)
        transcript = result["text"]
        st.subheader("Transcripción:")
        st.write(transcript)

        txt_bytes = transcript.encode("utf-8")
        st.download_button("📄 Descargar como .txt", data=txt_bytes, file_name="transcripcion.txt", mime="text/plain")

        doc = Document()
        doc.add_heading("Transcripción de Audio", level=1)
        doc.add_paragraph(transcript)
        docx_io = BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        st.download_button("📝 Descargar como .docx", data=docx_io, file_name="transcripcion.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    except Exception as e:
        st.error(f"Error durante la transcripción: {e}")
    finally:
        os.remove(tmp_path)
